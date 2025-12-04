from __future__ import annotations

import csv
import json
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional, Sequence

DEFAULT_CHUNK_SIZE = 512
DEFAULT_CHUNK_OVERLAP = 80


def _default_tokenizer(text: str) -> List[str]:
    """Very lightweight tokenizer: whitespace split with basic cleanup."""
    return [token for token in text.replace("\n", " ").split(" ") if token.strip()]


@dataclass
class ChunkingConfig:
    """Configuration controlling text chunking."""

    chunk_size: int = DEFAULT_CHUNK_SIZE
    overlap: int = DEFAULT_CHUNK_OVERLAP
    tokenizer: Callable[[str], List[str]] = _default_tokenizer


@dataclass
class Document:
    """Unified document representation used across ingestion and retrieval."""

    id: str
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    embedding: Optional[List[float]] = None

    @property
    def doc_id(self) -> str:
        # Backward-compatible alias for legacy retriever/tests.
        return self.id

    @property
    def source(self) -> str:
        return str(self.metadata.get("source", ""))


DocumentChunk = Document  # alias for backward compatibility


def chunk_text(text: str, config: ChunkingConfig | None = None) -> List[str]:
    """
    Splits text into overlapping chunks.

    Args:
        text: raw text to split.
        config: chunking configuration (size in tokens and overlap).
    """
    cfg = config or ChunkingConfig()
    tokens = cfg.tokenizer(text)
    if not tokens:
        return []

    chunks: List[str] = []
    step = max(1, cfg.chunk_size - cfg.overlap)
    for start in range(0, len(tokens), step):
        window = tokens[start : start + cfg.chunk_size]
        if window:
            chunks.append(" ".join(window))
    return chunks


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _read_pdf_file(path: Path) -> str:
    """Best-effort PDF reader; falls back to empty string if no deps."""
    try:
        import PyPDF2  # type: ignore
    except Exception:
        return ""
    text_parts: List[str] = []
    try:
        with path.open("rb") as stream:
            reader = PyPDF2.PdfReader(stream)
            for page in reader.pages:
                extracted = page.extract_text() or ""
                text_parts.append(extracted)
    except Exception:
        return ""
    return "\n".join(text_parts)


def _read_docx_file(path: Path) -> str:
    """Reads DOCX with optional dependency."""
    try:
        import docx  # type: ignore
    except Exception:
        return ""
    try:
        document = docx.Document(str(path))
        return "\n".join([para.text for para in document.paragraphs])
    except Exception:
        return ""


def _read_price_csv(path: Path) -> List[Dict[str, Any]]:
    with path.open("r", newline="", encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        if reader.fieldnames:
            return [dict(row) for row in reader]
        handle.seek(0)
        reader_generic = csv.reader(handle)
        return [{"row": " ".join(row)} for row in reader_generic]


def _read_price_excel(path: Path) -> List[Dict[str, Any]]:
    """Reads Excel via pandas if available; returns list of row dicts."""
    try:
        import pandas as pd  # type: ignore
    except Exception:
        return []
    try:
        df = pd.read_excel(path)
    except Exception:
        return []
    return df.fillna("").to_dict(orient="records")


def _build_doc_id(path: Path, idx: int) -> str:
    return f"{path.stem}:{idx}"


def _base_metadata(
    path: Path, source_type: str, extra: Dict[str, Any] | None = None
) -> Dict[str, Any]:
    meta: Dict[str, Any] = {
        "source": str(path),
        "source_type": source_type,
        "lang": None,
        "ingested_at": datetime.utcnow().isoformat(),
    }
    if extra:
        meta.update(extra)
    return meta


def _documents_from_rows(
    rows: Iterable[Dict[str, Any]],
    path: Path,
    source_type: str,
    chunking: ChunkingConfig,
    extra_metadata: Dict[str, Any] | None = None,
) -> List[Document]:
    docs: List[Document] = []
    for idx, row in enumerate(rows):
        text = " ".join(f"{k}: {v}" for k, v in row.items() if v not in (None, ""))
        for chunk_idx, chunk in enumerate(chunk_text(text, chunking)):
            doc_id = f"{_build_doc_id(path, idx)}#{chunk_idx}"
            metadata = _base_metadata(path, source_type, extra_metadata)
            metadata.update({"row_index": idx})
            docs.append(Document(id=doc_id, text=chunk, metadata=metadata))
    return docs


def load_price_documents(
    path: str | Path,
    chunking: ChunkingConfig | None = None,
    base_metadata: Dict[str, Any] | None = None,
) -> List[Document]:
    """Loads price lists from CSV/Excel and returns chunked documents."""
    cfg = chunking or ChunkingConfig()
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    rows: List[Dict[str, Any]]
    if suffix in {".csv"}:
        rows = _read_price_csv(file_path)
    elif suffix in {".xlsx", ".xls"}:
        rows = _read_price_excel(file_path)
    else:
        return []
    return _documents_from_rows(rows, file_path, "price_list", cfg, base_metadata)


def load_contract_documents(
    path: str | Path,
    chunking: ChunkingConfig | None = None,
    base_metadata: Dict[str, Any] | None = None,
) -> List[Document]:
    """Loads contract/regulation docs (txt/pdf/docx) and chunks them."""
    cfg = chunking or ChunkingConfig()
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    text = ""
    if suffix == ".txt":
        text = _read_text_file(file_path)
    elif suffix == ".pdf":
        text = _read_pdf_file(file_path)
    elif suffix in {".docx", ".doc"}:
        text = _read_docx_file(file_path)
    if not text:
        return []
    docs: List[Document] = []
    for idx, chunk in enumerate(chunk_text(text, cfg)):
        metadata = _base_metadata(file_path, "contract", base_metadata)
        metadata["page_chunk"] = idx
        docs.append(Document(id=_build_doc_id(file_path, idx), text=chunk, metadata=metadata))
    return docs


def load_messages_documents(
    messages: Sequence[Dict[str, Any]],
    source_name: str = "chatlog",
    chunking: ChunkingConfig | None = None,
    base_metadata: Dict[str, Any] | None = None,
) -> List[Document]:
    """
    Converts chat/email logs into documents.

    Expected message dict fields: `id`, `text`, `author`, `client_id`, `timestamp`.
    """
    cfg = chunking or ChunkingConfig()
    docs: List[Document] = []
    for idx, msg in enumerate(messages):
        text = msg.get("text") or ""
        if not text:
            continue
        meta = {
            "source_type": "message",
            "source": source_name,
            "author": msg.get("author"),
            "client_id": msg.get("client_id"),
            "timestamp": msg.get("timestamp"),
        }
        if base_metadata:
            meta.update(base_metadata)
        for chunk_idx, chunk in enumerate(chunk_text(text, cfg)):
            doc_id = f"{source_name}:{msg.get('id', idx)}#{chunk_idx}"
            docs.append(Document(id=doc_id, text=chunk, metadata=meta))
    return docs


def ingest_from_dir(
    directory: str | Path,
    chunking: ChunkingConfig | None = None,
    default_lang: str | None = "ru",
) -> List[Document]:
    """
    Ingests mixed corpus from a directory.

    Supports .txt/.pdf/.docx files and CSV/Excel price lists.
    """
    cfg = chunking or ChunkingConfig()
    root = Path(directory)
    docs: List[Document] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        base_meta = {"lang": default_lang}
        if suffix in {".txt", ".pdf", ".docx", ".doc"}:
            docs.extend(load_contract_documents(path, cfg, base_meta))
        elif suffix in {".csv", ".xlsx", ".xls"}:
            docs.extend(load_price_documents(path, cfg, base_meta))
        elif suffix in {".json"} and path.name.startswith("messages"):
            try:
                messages = json.loads(path.read_text(encoding="utf-8"))
                if isinstance(messages, list):
                    docs.extend(load_messages_documents(messages, path.stem, cfg, base_meta))
            except Exception:
                continue
    return docs
